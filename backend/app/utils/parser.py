import logging

import pdfplumber
import docx

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        logger.error(f"PDF extraction error on '{file_path}': {e}")
        raise
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file using python-docx."""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction error on '{file_path}': {e}")
        raise
